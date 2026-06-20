"""Extracción de texto desde distintos tipos de documento."""
from pathlib import Path

PLAIN_TEXT_EXTS = {".txt", ".md", ".markdown", ".csv", ".json", ".log", ".rst", ".py", ".html", ".xml"}
SUPPORTED_EXTS = {".pdf", ".docx"} | PLAIN_TEXT_EXTS


class ExtractionError(Exception):
    pass


def is_supported(path) -> bool:
    return Path(path).suffix.lower() in SUPPORTED_EXTS


def extract_text(path) -> str:
    ext = Path(path).suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(path)
    if ext == ".docx":
        return _extract_docx(path)
    if ext in PLAIN_TEXT_EXTS:
        return _extract_plain(path)
    raise ExtractionError(f"Tipo de archivo no soportado: {ext}")


def _extract_pdf(path) -> str:
    from pypdf import PdfReader

    try:
        reader = PdfReader(str(path))
    except Exception as e:  # noqa: BLE001
        raise ExtractionError(f"No se pudo leer el PDF: {e}") from e
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001
            parts.append("")
    text = "\n".join(parts).strip()
    if not text:
        raise ExtractionError(
            "El PDF no contiene texto extraíble (puede ser un escaneo/imagen). "
            "Se necesitaría OCR, que esta app no incluye todavía."
        )
    return text


def _extract_docx(path) -> str:
    import docx

    try:
        document = docx.Document(str(path))
    except Exception as e:  # noqa: BLE001
        raise ExtractionError(f"No se pudo leer el DOCX: {e}") from e
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    # También el texto de las tablas.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts).strip()
    if not text:
        raise ExtractionError("El documento DOCX está vacío.")
    return text


def _extract_plain(path) -> str:
    raw = Path(path).read_bytes()
    for enc in ("utf-8", "latin-1"):
        try:
            text = raw.decode(enc).strip()
            if text:
                return text
        except UnicodeDecodeError:
            continue
    raise ExtractionError("No se pudo decodificar el archivo de texto.")
